from dataclasses import dataclass
from pathlib import Path
from typing import List


@dataclass
class LoadedDocument:
    text: str
    source_file: str
    source_type: str


def load_txt_documents(data_dir: str) -> List[LoadedDocument]:
    base = Path(data_dir)
    docs: List[LoadedDocument] = []
    for p in base.rglob("*.txt"):
        text = p.read_text(encoding="utf-8", errors="ignore").strip()
        if text:
            docs.append(LoadedDocument(text=text, source_file=str(p), source_type="txt"))
    return docs


def load_pdf_documents(data_dir: str) -> List[LoadedDocument]:
    try:
        from pypdf import PdfReader
    except Exception:
        return []

    base = Path(data_dir)
    docs: List[LoadedDocument] = []
    for p in base.rglob("*.pdf"):
        try:
            reader = PdfReader(str(p))
            pages = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    pages.append(t.strip())
            if pages:
                docs.append(LoadedDocument(text="\n".join(pages), source_file=str(p), source_type="pdf"))
        except Exception:
            continue
    return docs


def load_docx_documents(data_dir: str) -> List[LoadedDocument]:
    try:
        from docx import Document
    except Exception:
        return []

    base = Path(data_dir)
    docs: List[LoadedDocument] = []
    for p in base.rglob("*.docx"):
        try:
            doc = Document(str(p))
            paragraphs = [x.text.strip() for x in doc.paragraphs if x.text and x.text.strip()]
            if paragraphs:
                docs.append(LoadedDocument(text="\n".join(paragraphs), source_file=str(p), source_type="docx"))
        except Exception:
            continue
    return docs


def load_course_documents(data_dir: str) -> List[LoadedDocument]:
    docs: List[LoadedDocument] = []
    docs.extend(load_txt_documents(data_dir))
    docs.extend(load_pdf_documents(data_dir))
    docs.extend(load_docx_documents(data_dir))
    return docs